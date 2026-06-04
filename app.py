import os
import re
import json
import time
import uuid
import secrets
import subprocess
import sys
import socket as _socket
import threading
from flask import (Flask, request, jsonify, render_template,
                   make_response, session, redirect, url_for)

# ── Force IPv4 globally ───────────────────────────────────────────────────────
# Render (and some other Linux hosts) have IPv6 in DNS but no IPv6 routing.
# Python tries IPv6 first → OSError errno 101 (ENETUNREACH).
# This patch makes every library (requests, smtplib, httpx/openai) use IPv4.
_orig_gai = _socket.getaddrinfo
def _ipv4_only(host, port, family=0, type=0, proto=0, flags=0):
    return _orig_gai(host, port, _socket.AF_INET, type, proto, flags)
_socket.getaddrinfo = _ipv4_only

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))

BASE          = os.path.dirname(os.path.abspath(__file__))
_USERDATA_DIR = os.path.join(BASE, "user_data")
_EXTRACTS_DIR = os.path.join(_USERDATA_DIR, "_extracts")
os.makedirs(_EXTRACTS_DIR, exist_ok=True)

# Site password gate
_SITE_PASSWORD = os.environ.get("SITE_PASSWORD", "")

# ── Per-run state ─────────────────────────────────────────────────────────────
# Multiple users can run pipelines simultaneously.
# Each run gets a UUID; the client polls /api/status?run_id=<uuid>.
_active_runs  = {}            # run_id → {running, output, done, error}
_user_locks   = {}            # email  → Lock  (one pipeline per email at a time)
_state_lock   = threading.Lock()


def _is_authenticated():
    if not _SITE_PASSWORD:
        return True
    return session.get("auth") == _SITE_PASSWORD


def _user_dir(email: str) -> str:
    """Return (and create) the per-user data directory for *email*."""
    safe = re.sub(r"[^a-zA-Z0-9._-]", "_", email.lower().strip())
    path = os.path.join(_USERDATA_DIR, safe)
    os.makedirs(os.path.join(path, "output", "resumes"),       exist_ok=True)
    os.makedirs(os.path.join(path, "output", "cover_letters"), exist_ok=True)
    return path


def _get_user_lock(email: str) -> threading.Lock:
    with _state_lock:
        if email not in _user_locks:
            _user_locks[email] = threading.Lock()
        return _user_locks[email]


def _extract_path() -> str:
    """Per-session path for the last successfully extracted resume text."""
    key = session.get("extract_key")
    if not key:
        key = secrets.token_hex(16)
        session["extract_key"] = key
    path = os.path.join(_EXTRACTS_DIR, f"{key}.txt")
    return path


# ── Login ─────────────────────────────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    error = ""
    if request.method == "POST":
        if request.form.get("password") == _SITE_PASSWORD:
            session["auth"] = _SITE_PASSWORD
            return redirect(url_for("index"))
        error = "Wrong password — try again."
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>JobMatchAI — Login</title>
  <style>
    *{{box-sizing:border-box;margin:0;padding:0}}
    body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
         background:#0f172a;display:flex;align-items:center;
         justify-content:center;min-height:100vh}}
    .card{{background:#1e293b;border-radius:16px;padding:40px 36px;
           width:100%;max-width:360px;box-shadow:0 20px 60px rgba(0,0,0,.5)}}
    h1{{color:#38bdf8;font-size:1.5rem;margin-bottom:6px;text-align:center}}
    p{{color:#94a3b8;font-size:.85rem;text-align:center;margin-bottom:28px}}
    input{{width:100%;padding:12px 14px;border-radius:8px;border:1.5px solid #334155;
           background:#0f172a;color:#f1f5f9;font-size:1rem;outline:none;
           transition:border .2s}}
    input:focus{{border-color:#38bdf8}}
    button{{margin-top:16px;width:100%;padding:13px;border-radius:8px;border:none;
            background:linear-gradient(135deg,#2563eb,#0ea5e9);color:#fff;
            font-size:1rem;font-weight:600;cursor:pointer;transition:opacity .2s}}
    button:hover{{opacity:.88}}
    .err{{color:#f87171;font-size:.84rem;margin-top:12px;text-align:center}}
  </style>
</head>
<body>
  <div class="card">
    <h1>🤖 JobMatchAI</h1>
    <p>Enter the site password to continue</p>
    <form method="post">
      <input type="password" name="password" placeholder="Password" autofocus required>
      <button type="submit">Enter</button>
    </form>
    {'<p class="err">' + error + '</p>' if error else ''}
  </div>
</body>
</html>"""


# ── Main page ─────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    if not _is_authenticated():
        return redirect(url_for("login"))
    resp = make_response(render_template("index.html"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp


# ── Run pipeline ──────────────────────────────────────────────────────────────

@app.route("/api/run", methods=["POST"])
def run_pipeline():
    if not _is_authenticated():
        return jsonify({"error": "Not authenticated. Please log in first."}), 401

    data = request.get_json(force=True, silent=True) or {}

    # Validate required fields
    if not data.get("name") or not data.get("email"):
        return jsonify({"error": "Name and email are required."}), 400
    if not data.get("target_roles"):
        return jsonify({"error": "At least one target role is required."}), 400

    email      = data["email"].strip()
    user_lock  = _get_user_lock(email)

    if not user_lock.acquire(blocking=False):
        return jsonify({"error": "A pipeline is already running for this account. Please wait."}), 409

    try:
        resume_txt = data.get("resume_text", "").strip()
        if not resume_txt:
            user_lock.release()
            return jsonify({"error": "Resume text is required."}), 400

        # ── Short-text auto-recovery ─────────────────────────────────────────
        if len(resume_txt) < 1500:
            recovered = False
            try:
                cache_path = _extract_path()
                cache_age  = time.time() - os.path.getmtime(cache_path)
                if cache_age < 1800:
                    with open(cache_path, "r", encoding="utf-8") as f:
                        cached = f.read().strip()
                    if len(cached) > len(resume_txt):
                        resume_txt = cached
                        recovered  = True
            except Exception:
                pass

            if not recovered:
                user_lock.release()
                return jsonify({
                    "error": (
                        f"Resume text appears incomplete "
                        f"({len(resume_txt)} characters — a full resume is usually 2 000+).\n"
                        "Please re-upload your resume file or paste the full text."
                    )
                }), 400

        config = {
            "name":                 data.get("name", "").strip(),
            "email":                email,
            "receiver_email":       data.get("receiver_email", "").strip() or email,
            "phone":                data.get("phone", "").strip(),
            "location":             data.get("location", "").strip().title(),
            "linkedin":             data.get("linkedin", "").strip(),
            "portfolio":            data.get("portfolio", "").strip(),
            "target_roles":         data.get("target_roles", []),
            "max_years_experience": int(data.get("max_years_experience", 0)),
            "match_threshold":      int(data.get("match_threshold", 70)),
            "max_jobs":             int(data.get("max_jobs", 15)),
            "job_sources":          data.get("job_sources", {
                "adzuna": True, "remotive": True, "arbeitnow": True,
                "greenhouse": True, "themuse": True, "remoteok": True,
                "jobicy": True, "usajobs": True, "jooble": True, "findwork": True,
                "indeed": True, "lever": True, "weworkremotely": True,
                "workingnomads": True, "hnhiring": True,
            }),
            "greenhouse_companies": data.get("greenhouse_companies", []),
            "lever_companies":      data.get("lever_companies", []),
        }

        udir = _user_dir(email)
        with open(os.path.join(udir, "config.json"), "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2)
        with open(os.path.join(udir, "resume.txt"), "w", encoding="utf-8") as f:
            f.write(resume_txt)

        run_id = uuid.uuid4().hex
        with _state_lock:
            _active_runs[run_id] = {
                "running": True, "output": [], "done": False, "error": None
            }

    except Exception as exc:
        user_lock.release()
        return jsonify({"error": str(exc)}), 500

    def _run():
        try:
            env          = os.environ.copy()
            env["USER_DIR"] = udir
            proc = subprocess.Popen(
                [sys.executable, "-W", "ignore", "main.py"],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
                cwd=BASE,
                env=env,
            )
            for line in proc.stdout:
                line = line.rstrip()
                if not line:
                    continue
                if "site-packages" in line or "warnings.warn" in line:
                    continue
                if line.strip().startswith("C:\\") or line.strip().startswith("/usr/"):
                    continue
                _active_runs[run_id]["output"].append(line)
            proc.wait()
            if proc.returncode != 0:
                _active_runs[run_id]["error"] = f"Process exited with code {proc.returncode}"
        except Exception as exc:
            _active_runs[run_id]["error"] = str(exc)
        finally:
            _active_runs[run_id]["running"] = False
            _active_runs[run_id]["done"]    = True
            _active_runs[run_id]["_ts"]     = time.time()
            user_lock.release()
            _prune_old_runs()

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"message": "Pipeline started", "run_id": run_id})


def _prune_old_runs():
    """Remove completed runs from memory after 30 minutes."""
    cutoff = time.time() - 1800
    with _state_lock:
        stale = [rid for rid, s in _active_runs.items()
                 if s.get("done") and s.get("_ts", 0) < cutoff]
        for rid in stale:
            del _active_runs[rid]


# ── Parse resume ──────────────────────────────────────────────────────────────

@app.route("/api/parse-resume", methods=["POST"])
def parse_resume():
    if not _is_authenticated():
        return jsonify({"error": "Not authenticated."}), 401

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    fname = file.filename.lower()

    def _save_cache(text):
        try:
            with open(_extract_path(), "w", encoding="utf-8") as f:
                f.write(text)
        except Exception:
            pass

    # ── Plain text ────────────────────────────────────────────────────────
    if fname.endswith(".txt"):
        try:
            text = file.read().decode("utf-8", errors="replace")
            _save_cache(text)
            return jsonify({"text": text})
        except Exception as e:
            return jsonify({"error": str(e)}), 400

    # ── PDF ───────────────────────────────────────────────────────────────
    if fname.endswith(".pdf"):
        file_bytes = file.read()

        def _clean_pdf_text(raw):
            raw = raw.replace("\t", " ")
            raw = re.sub(r"[ ]{2,}", " ", raw)
            lines = [l.rstrip() for l in raw.splitlines()]
            return "\n".join(lines).strip()

        try:
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(file_bytes), strict=False)
            pages  = [p.extract_text() or "" for p in reader.pages]
            text   = _clean_pdf_text("\n\n".join(pages))
            if text:
                _save_cache(text)
                return jsonify({"text": text})
        except Exception:
            pass

        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = _clean_pdf_text("\n\n".join(pages))
            if text:
                _save_cache(text)
                return jsonify({"text": text})
        except Exception:
            pass

        return jsonify({
            "error": "No text could be extracted — the PDF may be scanned/image-based. "
                     "Please paste the text manually."
        }), 400

    # ── DOCX ──────────────────────────────────────────────────────────────
    if fname.endswith(".docx"):
        try:
            from docx import Document
            from docx.oxml.ns import qn
            import io
            raw = file.read()
            doc = Document(io.BytesIO(raw))
            lines = []
            seen  = set()

            def _add(t):
                t = t.strip()
                if t and t not in seen:
                    lines.append(t)
                    seen.add(t)

            for p in doc.paragraphs:
                _add(p.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _add(para.text)
            try:
                for txbx in doc.element.body.iter(qn("w:txbxContent")):
                    for para in txbx.iter(qn("w:p")):
                        t = "".join(run.text for run in para.iter(qn("w:t")))
                        _add(t)
            except Exception:
                pass
            try:
                for section in doc.sections:
                    for hdr in (section.header, section.footer):
                        if hdr and not hdr.is_linked_to_previous:
                            for p in hdr.paragraphs:
                                _add(p.text)
            except Exception:
                pass

            text = "\n".join(lines).strip()
            if not text:
                return jsonify({"error": "No text found in the DOCX file."}), 400
            _save_cache(text)
            return jsonify({"text": text})
        except ImportError:
            return jsonify({"error": "python-docx is not installed."}), 500
        except Exception as e:
            return jsonify({"error": f"DOCX parsing error: {e}"}), 400

    # ── DOC ───────────────────────────────────────────────────────────────
    if fname.endswith(".doc"):
        file_bytes = file.read()

        try:
            import win32com.client
            import tempfile, pythoncom
            pythoncom.CoInitialize()
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
                tf.write(file_bytes)
                tmp_path = tf.name
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                doc  = word.Documents.Open(os.path.abspath(tmp_path),
                                           ReadOnly=True, AddToRecentFiles=False)
                text = doc.Content.Text.strip()
                doc.Close(False)
                word.Quit()
            finally:
                try: os.unlink(tmp_path)
                except: pass
                try: pythoncom.CoUninitialize()
                except: pass
            if text:
                _save_cache(text)
                return jsonify({"text": text})
        except ImportError:
            pass
        except Exception:
            pass

        try:
            from docx import Document
            import io
            doc   = Document(io.BytesIO(file_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            seen  = set(lines)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            t = para.text.strip()
                            if t and t not in seen:
                                lines.append(t)
                                seen.add(t)
            text = "\n".join(lines).strip()
            if text:
                _save_cache(text)
                return jsonify({"text": text})
        except Exception:
            pass

        try:
            raw_text = file_bytes.decode("utf-16-le", errors="ignore")
            salvaged = re.sub(r"[^\x20-\x7E\xA0-\xFF\n\t]+", " ", raw_text)
            salvaged = re.sub(r" {3,}", " ", salvaged).strip()
            if len(salvaged) > 200:
                _save_cache(salvaged)
                return jsonify({"text": salvaged})
        except Exception:
            pass

        return jsonify({
            "error": (
                "Could not extract text from this .doc file.\n"
                "Save as .docx in Word or Google Docs and re-upload."
            )
        }), 400

    return jsonify({
        "error": "Unsupported format. Please upload .txt, .pdf, .docx, or .doc."
    }), 400


# ── Status polling ────────────────────────────────────────────────────────────

@app.route("/api/status")
def get_status():
    if not _is_authenticated():
        return jsonify({"error": "Not authenticated."}), 401
    run_id = request.args.get("run_id", "")
    with _state_lock:
        status = _active_runs.get(run_id)
    if status is None:
        return jsonify({"error": "Run not found", "done": True}), 404
    return jsonify(status)


if __name__ == "__main__":
    print("JobMatchAI for Everyone — starting at http://localhost:5000")
    app.run(host="0.0.0.0", port=5000, debug=False)
